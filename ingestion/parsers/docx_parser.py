import uuid
from pathlib import Path
from core.models.document import Document
from core.models.section import Section
from core.models.block import Block
from .base import BaseParser
from docx import Document as DocxDocument


class DocxParser(BaseParser):
    def supports(self, file_path: Path) -> bool:
        return file_path.suffix.lower() in [".docx", ".doc"]

    def parse(self, file_path: Path) -> Document:
        docx = DocxDocument(file_path)

        paragraphs = [p.text for p in docx.paragraphs if p.text.strip()]
        content = "\n\n".join(paragraphs)

        block = Block(
            block_id=str(uuid.uuid4()), type="text", content=content, metadata={}
        )

        section = Section(
            section_id=str(uuid.uuid4()),
            title=file_path.stem,
            level=1,
            blocks=[block],
            metadata={},
        )

        document = Document(
            document_id=str(uuid.uuid4()),
            sections=[section],
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx",
            },
        )

        return document
